from fastapi import FastAPI, HTTPException, BackgroundTasks, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import io
import json
import pandas as pd
import asyncio
from contextlib import asynccontextmanager
import uvicorn
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from agent import MongoDBAgent, phoenix_span_manager
from traces.setup import EvaluationPipeline
from traces.upload_dataset import PhoenixDatasetUploader
from websocket_handler import handle_chat_websocket, ws_manager
from qdrant.qdrant_initializer import RAGTool
from tools import filter_and_transform_content
from planner import plan_and_execute_query
# Pydantic models for API requests/responses
class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None

class ChatResponse(BaseModel):
    response: str
    conversation_id: str
    timestamp: str

class ToolCall(BaseModel):
    name: str
    arguments: Dict[str, Any]
    result: Optional[str] = None

class Message(BaseModel):
    id: str
    type: str  # "user", "assistant", "tool", "thought"
    content: str
    timestamp: str
    tool_name: Optional[str] = None
    tool_output: Optional[Any] = None


class ExportRequest(BaseModel):
    """Generic export request payload.

    Use one of:
    - rows: list of dicts to export directly
    - tool + query: server will re-run the query and export results
    - content: for DOCX exports of plain text
    """
    tool: Optional[str] = None  # 'mongo_query' | 'rag_search'
    query: Optional[str] = None
    params: Optional[Dict[str, Any]] = None
    rows: Optional[List[Dict[str, Any]]] = None
    content: Optional[str] = None
    title: Optional[str] = None


async def _resolve_rows_for_export(req: ExportRequest) -> List[Dict[str, Any]]:
    """Return list of row dicts for export, re-running query if needed."""
    if req.rows and isinstance(req.rows, list):
        return req.rows  # type: ignore[return-value]

    # Re-run a Mongo query if provided
    if (req.tool or "").lower() == "mongo_query" and req.query:
        try:
            result = await plan_and_execute_query(req.query)
            if not result.get("success"):
                raise HTTPException(status_code=400, detail=f"Query failed: {result.get('error')}")

            raw = result.get("result")
            intent = result.get("intent") or {}
            primary_entity = intent.get("primary_entity") if isinstance(intent, dict) else None
            filtered = filter_and_transform_content(raw, primary_entity=primary_entity)
            if isinstance(filtered, dict):
                return [filtered]
            if isinstance(filtered, list):
                # Ensure only dict-like rows are returned
                return [x for x in filtered if isinstance(x, dict)]  # type: ignore[list-item]
            return []
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to generate rows: {e}")

    # If nothing resolvable
    raise HTTPException(status_code=400, detail="No rows or resolvable query provided for export")

# Global MongoDB agent instance
mongodb_agent = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage the lifespan of the FastAPI application"""
    global mongodb_agent

    # Startup
    print("Starting MongoDB Agent with Phoenix tracing...")
    await phoenix_span_manager.initialize()
    mongodb_agent = MongoDBAgent()
    await mongodb_agent.initialize_tracing()
    await mongodb_agent.connect()
    print("MongoDB Agent connected successfully!")
    await RAGTool.initialize()
    print("RAGTool initialized successfully!")
    yield

    # Shutdown
    print("Shutting down MongoDB Agent...")
    await mongodb_agent.disconnect()
    print("MongoDB Agent disconnected.")

# Create FastAPI app
app = FastAPI(
    title="PMS Assistant API",
    description="Project Management System Assistant with MongoDB integration",
    version="1.0.0",
    lifespan=lifespan
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    # Allow all origins to simplify local/dev usage across ports (8080, 5173, etc.)
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "PMS Assistant API", "status": "running"}

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy"}

@app.websocket("/ws/chat")
async def websocket_chat(websocket: WebSocket):
    """WebSocket endpoint for streaming chat"""
    global mongodb_agent

    # Initialize agent if not already done (for testing/development)
    if not mongodb_agent:
        print("Initializing MongoDB Agent for WebSocket...")
        await phoenix_span_manager.initialize()
        mongodb_agent = MongoDBAgent()
        await mongodb_agent.initialize_tracing()
        await mongodb_agent.connect()

    await handle_chat_websocket(websocket, mongodb_agent)


@app.post("/eval/run")
async def run_evaluation(sample_size: int | None = None):
    """Run the evaluation pipeline and return the report."""
    pipeline = EvaluationPipeline()
    await pipeline.initialize()
    results = await pipeline.run_evaluation(sample_size=sample_size)
    report = await pipeline.generate_evaluation_report(results)
    return {"report": report}


@app.post("/phoenix/dataset/upload")
async def upload_phoenix_dataset():
    """Build and save the Phoenix evaluation dataset JSON locally."""
    uploader = PhoenixDatasetUploader()
    dataset = uploader.load_test_dataset()
    if not dataset:
        raise HTTPException(status_code=400, detail="No dataset loaded")
    dataset_info = uploader.create_phoenix_dataset(dataset)
    if not dataset_info.get("success"):
        raise HTTPException(status_code=500, detail="Failed to create Phoenix dataset")
    success = uploader.upload_to_phoenix(dataset_info)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to save dataset JSON")
    return {"success": True, "entries": len(dataset)}


@app.post("/export/csv")
async def export_csv(req: ExportRequest):
    """Export provided rows or resolvable query results to CSV."""
    rows = await _resolve_rows_for_export(req)
    if not rows:
        raise HTTPException(status_code=400, detail="No data to export")

    df = pd.DataFrame(rows)
    csv_bytes = df.to_csv(index=False).encode("utf-8")
    filename = (req.title or "export").replace(" ", "_") + ".csv"
    return StreamingResponse(
        io.BytesIO(csv_bytes),
        media_type="text/csv",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        },
    )


@app.post("/export/xlsx")
async def export_xlsx(req: ExportRequest):
    """Export provided rows or resolvable query results to XLSX."""
    rows = await _resolve_rows_for_export(req)
    if not rows:
        raise HTTPException(status_code=400, detail="No data to export")

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df = pd.DataFrame(rows)
        # Limit very wide cells to avoid bloating workbook
        df = df.applymap(lambda v: str(v) if not isinstance(v, (int, float)) else v)
        df.to_excel(writer, index=False, sheet_name="Data")
    output.seek(0)
    filename = (req.title or "export").replace(" ", "_") + ".xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        },
    )


@app.post("/export/docx")
async def export_docx(req: ExportRequest):
    """Export content or rows to a Word document (.docx)."""
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx not available: {e}")

    document = Document()
    title = req.title or "Export"
    document.add_heading(title, level=1)

    # Prefer tabular data if available
    rows: List[Dict[str, Any]] = []
    if req.rows:
        rows = req.rows  # type: ignore[assignment]
    elif (req.tool or "").lower() == "mongo_query" and req.query:
        rows = await _resolve_rows_for_export(req)

    if rows:
        # Build table with columns from union of keys
        all_keys: List[str] = []
        seen = set()
        for r in rows:
            if isinstance(r, dict):
                for k in r.keys():
                    if k not in seen:
                        seen.add(k)
                        all_keys.append(k)
        if not all_keys:
            all_keys = ["value"]

        table = document.add_table(rows=1, cols=len(all_keys))
        hdr_cells = table.rows[0].cells
        for i, k in enumerate(all_keys):
            hdr_cells[i].text = str(k)
        for r in rows:
            row_cells = table.add_row().cells
            for i, k in enumerate(all_keys):
                val = r.get(k, "") if isinstance(r, dict) else ""
                row_cells[i].text = str(val)
    elif req.content:
        # Plain content paragraphs
        for para in str(req.content).split("\n\n"):
            document.add_paragraph(para)
    else:
        raise HTTPException(status_code=400, detail="No content or data to export")

    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0)

    filename = (title or "export").replace(" ", "_") + ".docx"
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        },
    )

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
